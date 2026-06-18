from __future__ import annotations

from pathlib import Path

from finalreview.models import DocumentChunk
from finalreview.parsers.base import BaseParser
from finalreview.utils.text_utils import normalize_text


class DOCXParser(BaseParser):
    file_types = {".docx"}

    def parse(self, file_path: Path) -> list[DocumentChunk]:
        try:
            from docx import Document
        except Exception as exc:
            return [
                self.make_chunk(
                    file_path,
                    f"DOCX 解析依赖 python-docx 不可用：{exc}",
                    0,
                    metadata={"parse_error": "missing_python_docx"},
                )
            ]
        try:
            doc = Document(str(file_path))
            chunks: list[DocumentChunk] = []
            for idx, paragraph in enumerate(doc.paragraphs, start=1):
                text = normalize_text(paragraph.text)
                if text:
                    chunks.append(self.make_chunk(file_path, text, idx, paragraph_index=idx))
            table_index = 0
            for table in doc.tables:
                table_index += 1
                rows = []
                for row in table.rows:
                    rows.append(" | ".join(cell.text.strip() for cell in row.cells))
                text = normalize_text("\n".join(rows))
                if text:
                    chunks.append(
                        self.make_chunk(
                            file_path,
                            text,
                            10_000 + table_index,
                            metadata={"table_index": table_index},
                        )
                    )
            return chunks or [self.make_chunk(file_path, "DOCX 未提取到文本。", 0)]
        except Exception as exc:
            return [
                self.make_chunk(
                    file_path,
                    f"DOCX 解析失败：{exc}",
                    0,
                    metadata={"parse_error": str(exc)},
                )
            ]
